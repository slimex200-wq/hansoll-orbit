from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import os
import re
import sqlite3
import sys
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree


STYLE_RE = re.compile(r"(?<!\d)(\d{9}[A-Z]?)(?!\d)", re.IGNORECASE)
SUPPORTED_EXTENSIONS = {
    ".csv",
    ".docx",
    ".eml",
    ".html",
    ".htm",
    ".md",
    ".pdf",
    ".pptx",
    ".txt",
    ".xlsm",
    ".xlsx",
}
MAX_TEXT_CHARS = 400_000
MAX_PDF_PAGES = 20
MAX_FILE_BYTES = 80 * 1024 * 1024


@dataclass(frozen=True)
class StyleHit:
    style_no: str
    location: str
    snippet: str
    source: str


def utc_now() -> str:
    return datetime.now(UTC).isoformat()


def normalize_text(value: object) -> str:
    text = str(value or "")
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def snippet_hash(snippet: str) -> str:
    return hashlib.sha256(snippet.encode("utf-8", "ignore")).hexdigest()[:16]


def find_styles(text: str) -> list[str]:
    seen: set[str] = set()
    styles: list[str] = []
    for match in STYLE_RE.finditer(text):
        style = match.group(1).upper()
        if style not in seen:
            seen.add(style)
            styles.append(style)
    return styles


def make_hits(text: str, location: str, source: str) -> list[StyleHit]:
    clean = normalize_text(text)
    if not clean:
        return []
    if len(clean) > 500:
        clean = clean[:500]
    return [StyleHit(style, location, clean, source) for style in find_styles(clean)]


def cap_hits(hits: list[StyleHit], max_per_style_file: int) -> list[StyleHit]:
    capped: list[StyleHit] = []
    counts: dict[str, int] = {}
    for hit in hits:
        count = counts.get(hit.style_no, 0)
        if count >= max_per_style_file:
            continue
        counts[hit.style_no] = count + 1
        capped.append(hit)
    return capped


def extract_xlsx(path: Path) -> list[StyleHit]:
    from openpyxl import load_workbook

    hits: list[StyleHit] = []
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        chars = 0
        for sheet in workbook.worksheets:
            for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                parts = [normalize_text(cell) for cell in row if cell is not None]
                if not parts:
                    continue
                row_text = " | ".join(part for part in parts if part)
                chars += len(row_text)
                if STYLE_RE.search(row_text):
                    hits.extend(make_hits(row_text, f"{sheet.title}!R{row_no}", "cell"))
                if chars >= MAX_TEXT_CHARS:
                    return hits
    finally:
        workbook.close()
    return hits


def extract_pdf(path: Path) -> list[StyleHit]:
    from pypdf import PdfReader

    hits: list[StyleHit] = []
    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if STYLE_RE.search(text):
            hits.extend(make_hits(text, f"page {index}", "pdf_text"))
    return hits


def extract_docx(path: Path) -> list[StyleHit]:
    from docx import Document

    hits: list[StyleHit] = []
    doc = Document(str(path))
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text or ""
        if STYLE_RE.search(text):
            hits.extend(make_hits(text, f"paragraph {idx}", "docx"))
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            text = " | ".join(normalize_text(cell.text) for cell in row.cells)
            if STYLE_RE.search(text):
                hits.extend(make_hits(text, f"table {table_idx} row {row_idx}", "docx"))
    return hits


def extract_pptx(path: Path) -> list[StyleHit]:
    hits: list[StyleHit] = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(name for name in archive.namelist() if name.startswith("ppt/slides/") and name.endswith(".xml"))
        for name in names:
            try:
                root = ElementTree.fromstring(archive.read(name))
            except Exception:
                continue
            texts: list[str] = []
            for element in root.iter():
                if element.text and element.text.strip():
                    texts.append(element.text.strip())
            slide_text = " ".join(texts)
            if STYLE_RE.search(slide_text):
                hits.extend(make_hits(slide_text, Path(name).stem, "pptx"))
    return hits


def extract_text_file(path: Path) -> list[StyleHit]:
    hits: list[StyleHit] = []
    try:
        chars = 0
        with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
            if path.suffix.lower() == ".csv":
                reader = csv.reader(handle)
                for row_no, row in enumerate(reader, start=1):
                    row_text = " | ".join(normalize_text(cell) for cell in row)
                    chars += len(row_text)
                    if STYLE_RE.search(row_text):
                        hits.extend(make_hits(row_text, f"line {row_no}", "text"))
                    if chars > MAX_TEXT_CHARS:
                        break
            else:
                for line_no, line in enumerate(handle, start=1):
                    chars += len(line)
                    if STYLE_RE.search(line):
                        hits.extend(make_hits(line, f"line {line_no}", "text"))
                    if chars > MAX_TEXT_CHARS:
                        break
    except OSError:
        raise
    return hits


def extract_hits(path: Path) -> tuple[str, list[StyleHit], str | None]:
    hits = make_hits(str(path), "path", "path")
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return "path_only", hits, None
    try:
        stat = path.stat()
        if stat.st_size > MAX_FILE_BYTES:
            return "skipped_large", hits, f"larger than {MAX_FILE_BYTES} bytes"
        if suffix in {".xlsx", ".xlsm"}:
            hits.extend(extract_xlsx(path))
        elif suffix == ".pdf":
            hits.extend(extract_pdf(path))
        elif suffix == ".docx":
            hits.extend(extract_docx(path))
        elif suffix == ".pptx":
            hits.extend(extract_pptx(path))
        elif suffix in {".csv", ".txt", ".md", ".html", ".htm", ".eml"}:
            hits.extend(extract_text_file(path))
        return "parsed", hits, None
    except Exception as exc:
        return "error", hits, f"{type(exc).__name__}: {exc}"


def init_db(db_path: Path, with_fts: bool) -> None:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(db_path) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS files (
                path TEXT PRIMARY KEY,
                relative_path TEXT NOT NULL,
                top_folder TEXT NOT NULL,
                extension TEXT NOT NULL,
                size INTEGER NOT NULL,
                mtime_ns INTEGER NOT NULL,
                parse_status TEXT NOT NULL,
                hit_count INTEGER NOT NULL,
                error TEXT,
                indexed_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS style_hits (
                style_no TEXT NOT NULL,
                path TEXT NOT NULL,
                relative_path TEXT NOT NULL,
                top_folder TEXT NOT NULL,
                extension TEXT NOT NULL,
                location TEXT NOT NULL,
                snippet TEXT NOT NULL,
                snippet_hash TEXT NOT NULL,
                source TEXT NOT NULL,
                indexed_at TEXT NOT NULL,
                PRIMARY KEY (style_no, path, location, snippet_hash)
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_style_hits_style ON style_hits(style_no)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_style_hits_path ON style_hits(path)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_files_top_folder ON files(top_folder)")
        if with_fts:
            conn.execute(
                """
                CREATE VIRTUAL TABLE IF NOT EXISTS style_hits_fts
                USING fts5(style_no, relative_path, top_folder, location, snippet)
                """
            )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS ingest_runs (
                run_id TEXT PRIMARY KEY,
                started_at TEXT NOT NULL,
                completed_at TEXT,
                stats_json TEXT NOT NULL
            )
            """
        )


def iter_files(root: Path, include_tops: set[str], path_contains: list[str]) -> Iterable[Path]:
    lowered_terms = [term.lower() for term in path_contains if term]
    for top in sorted(include_tops):
        start = root / top
        if not start.exists():
            continue
        for dirpath, dirnames, filenames in os.walk(start):
            dirnames[:] = [name for name in dirnames if not name.startswith(".")]
            for filename in filenames:
                if filename.startswith("~$"):
                    continue
                path = Path(dirpath) / filename
                if lowered_terms and not any(term in str(path).lower() for term in lowered_terms):
                    continue
                yield path


def index_file(
    conn: sqlite3.Connection,
    root: Path,
    path: Path,
    indexed_at: str,
    force: bool,
    with_fts: bool,
    max_hits_per_style_file: int,
) -> tuple[int, int, str]:
    relative_path = str(path.relative_to(root))
    top_folder = Path(relative_path).parts[0] if Path(relative_path).parts else ""
    suffix = path.suffix.lower()
    try:
        stat = path.stat()
    except OSError as exc:
        return 0, 1, f"stat_error: {exc}"

    existing = conn.execute(
        "SELECT size, mtime_ns FROM files WHERE path = ?",
        (str(path),),
    ).fetchone()
    if not force and existing and existing[0] == stat.st_size and existing[1] == stat.st_mtime_ns:
        return 0, 0, "unchanged"

    status, hits, error = extract_hits(path)
    hits = cap_hits(hits, max_hits_per_style_file)
    conn.execute("DELETE FROM style_hits WHERE path = ?", (str(path),))
    if with_fts:
        conn.execute("DELETE FROM style_hits_fts WHERE relative_path = ?", (relative_path,))
    for hit in hits:
        digest = snippet_hash(hit.snippet)
        conn.execute(
            """
            INSERT OR IGNORE INTO style_hits (
                style_no, path, relative_path, top_folder, extension, location,
                snippet, snippet_hash, source, indexed_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                hit.style_no,
                str(path),
                relative_path,
                top_folder,
                suffix,
                hit.location,
                hit.snippet,
                digest,
                hit.source,
                indexed_at,
            ),
        )
        if with_fts:
            conn.execute(
                """
                INSERT INTO style_hits_fts(style_no, relative_path, top_folder, location, snippet)
                VALUES (?, ?, ?, ?, ?)
                """,
                (hit.style_no, relative_path, top_folder, hit.location, hit.snippet),
            )

    conn.execute(
        """
        INSERT INTO files (
            path, relative_path, top_folder, extension, size, mtime_ns,
            parse_status, hit_count, error, indexed_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(path) DO UPDATE SET
            relative_path=excluded.relative_path,
            top_folder=excluded.top_folder,
            extension=excluded.extension,
            size=excluded.size,
            mtime_ns=excluded.mtime_ns,
            parse_status=excluded.parse_status,
            hit_count=excluded.hit_count,
            error=excluded.error,
            indexed_at=excluded.indexed_at
        """,
        (
            str(path),
            relative_path,
            top_folder,
            suffix,
            stat.st_size,
            stat.st_mtime_ns,
            status,
            len(hits),
            error,
            indexed_at,
        ),
    )
    return 1, 0, status


def build_index(args: argparse.Namespace) -> int:

    root = args.root.expanduser().resolve()
    include_tops = {item for item in args.include_top if item}
    if args.reset and args.db.exists():
        args.db.unlink()
    init_db(args.db, args.with_fts)

    started_at = utc_now()
    run_id = "business-style-" + datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    stats = {
        "root": str(root),
        "include_tops": sorted(include_tops),
        "files_seen": 0,
        "files_indexed": 0,
        "files_error": 0,
        "unchanged": 0,
        "by_status": {},
    }

    with sqlite3.connect(args.db) as conn:
        conn.execute(
            "INSERT OR REPLACE INTO ingest_runs(run_id, started_at, completed_at, stats_json) VALUES (?, ?, ?, ?)",
            (run_id, started_at, None, json.dumps(stats, ensure_ascii=False)),
        )
        for path in iter_files(root, include_tops, args.path_contains or []):
            stats["files_seen"] += 1
            indexed, errored, status = index_file(
                conn,
                root,
                path,
                started_at,
                args.force,
                args.with_fts,
                args.max_hits_per_style_file,
            )
            stats["files_indexed"] += indexed
            stats["files_error"] += errored
            if status == "unchanged":
                stats["unchanged"] += 1
            stats["by_status"][status] = stats["by_status"].get(status, 0) + 1
            if stats["files_seen"] % args.progress_every == 0:
                conn.commit()
                print(json.dumps({k: stats[k] for k in ("files_seen", "files_indexed", "files_error", "unchanged")}, ensure_ascii=False), flush=True)
        completed_at = utc_now()
        conn.execute(
            "UPDATE ingest_runs SET completed_at = ?, stats_json = ? WHERE run_id = ?",
            (completed_at, json.dumps(stats, ensure_ascii=False), run_id),
        )
        conn.commit()

        total_hits = conn.execute("SELECT COUNT(*) FROM style_hits").fetchone()[0]
        total_styles = conn.execute("SELECT COUNT(DISTINCT style_no) FROM style_hits").fetchone()[0]

    print(json.dumps({"run_id": run_id, "total_hits": total_hits, "total_styles": total_styles, **stats}, ensure_ascii=False, indent=2))
    return 0


def search_index(args: argparse.Namespace) -> int:
    init_db(args.db, with_fts=False)
    query = args.query.strip()
    like = f"%{query}%"
    with sqlite3.connect(args.db) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            """
            SELECT style_no, relative_path, location, snippet, source, indexed_at
            FROM style_hits
            WHERE style_no LIKE ? OR relative_path LIKE ? OR snippet LIKE ?
            ORDER BY
                CASE WHEN style_no = ? THEN 0 ELSE 1 END,
                indexed_at DESC,
                relative_path ASC
            LIMIT ?
            """,
            (like, like, like, query.upper(), args.limit),
        ).fetchall()
    print(json.dumps([dict(row) for row in rows], ensure_ascii=False, indent=2))
    return 0


def index_stats(args: argparse.Namespace) -> int:
    init_db(args.db, with_fts=False)
    with sqlite3.connect(args.db) as conn:
        stats = {
            "db": str(args.db),
            "files": conn.execute("SELECT COUNT(*) FROM files").fetchone()[0],
            "style_hits": conn.execute("SELECT COUNT(*) FROM style_hits").fetchone()[0],
            "styles": conn.execute(
                "SELECT COUNT(DISTINCT style_no) FROM style_hits"
            ).fetchone()[0],
            "latest_indexed_at": conn.execute("SELECT MAX(indexed_at) FROM files").fetchone()[0],
        }
        by_status = conn.execute(
            """
            SELECT parse_status, COUNT(*) AS count
            FROM files
            GROUP BY parse_status
            ORDER BY count DESC
            """
        ).fetchall()
    stats["by_status"] = {row[0]: row[1] for row in by_status}
    print(json.dumps(stats, ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Build and query a compact style-to-file index for selected business folders."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    build = subparsers.add_parser("build", help="Build or refresh the style index")
    build.add_argument("--root", required=True, type=Path)
    build.add_argument("--db", required=True, type=Path)
    build.add_argument("--include-top", action="append", required=True)
    build.add_argument("--path-contains", action="append")
    build.add_argument("--force", action="store_true")
    build.add_argument("--reset", action="store_true")
    build.add_argument("--with-fts", action="store_true")
    build.add_argument("--max-hits-per-style-file", type=int, default=3)
    build.add_argument("--progress-every", type=int, default=250)

    search = subparsers.add_parser("search", help="Search style references in the index")
    search.add_argument("--db", required=True, type=Path)
    search.add_argument("--query", required=True)
    search.add_argument("--limit", type=int, default=20)

    stats = subparsers.add_parser("stats", help="Show index health and counts")
    stats.add_argument("--db", required=True, type=Path)
    return parser


def normalize_legacy_build_args(argv: list[str]) -> list[str]:
    if len(argv) <= 1:
        return argv
    if argv[1] in {"build", "search", "stats", "-h", "--help"}:
        return argv
    return [argv[0], "build", *argv[1:]]


def main() -> int:
    parser = build_parser()
    args = parser.parse_args(normalize_legacy_build_args(sys.argv)[1:])
    if args.command == "build":
        return build_index(args)
    if args.command == "search":
        return search_index(args)
    if args.command == "stats":
        return index_stats(args)
    return 1


if __name__ == "__main__":
    sys.exit(main())
